from docx import Document

# Une fonction qui cree un document Word et ajoute du texte
def create_word_document(file_path):
    doc = Document()
    doc.add_heading('Titre du Document', level=1)
    doc.add_paragraph('Ceci est un paragraphe dans le document Word.')
    # Ajoute un caractère sensible pour la détection
    doc.add_paragraph('Voici un caractère sensible : \u200B')  # Caractère invisible
    doc.save(file_path)


# Liste des caractères sensibles à détecter
SENSITIVE_CHARS = ['\u200B', '\u200C', '\u200D']

def detect_sensitive_characters(docx_path):
    doc = Document(docx_path)

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if any(char in text for char in SENSITIVE_CHARS):
                print("Attention : caractère sensible repéré")
                return  # On arrête après la première détection
    print("Aucun caractère sensible détecté.")

if __name__ == "__main__":
    # Remplace par ton chemin de fichier
    create_word_document("test.docx")
    choice = input("Voulez-vous détecter les caractères sensibles dans un document existant ? (oui/non) : ").strip().lower()
    if choice == 'oui':
        fichier = input("Entrez le chemin du fichier Word : ").strip()
        detect_sensitive_characters(fichier)
    else:
        detect_sensitive_characters("test.docx")
